"""
antigravity/ingest.py — File ingestion: convert uploaded files to corpus Markdown.

Supports: .pdf, .docx, .doc, .txt, .md

At upload time, also splits into chunks, classifies each chunk,
tags it with the target model mode, and registers everything in the DocIndex.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from antigravity.doc_index import DocIndex


def ingest_file(
    file_path: str | Path,
    corpus_dir: str | Path,
    doc_index: "DocIndex | None" = None,
) -> dict:
    """
    Convert a file to Markdown, save in corpus, and register in DocIndex.

    Returns dict with: {"doc_id": str, "filename": str, "chunks": int, "size": int}
    """
    file_path = Path(file_path)
    corpus_dir = Path(corpus_dir)
    corpus_dir.mkdir(parents=True, exist_ok=True)

    ext = file_path.suffix.lower()
    filename = file_path.stem

    # Extract text based on file type
    if ext == ".pdf":
        text = _extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        text = _extract_docx(file_path)
    elif ext in (".txt", ".md", ".markdown"):
        text = file_path.read_text(encoding="utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx, .txt, .md")

    if not text.strip():
        raise ValueError(f"No text could be extracted from {file_path.name}")

    # Convert to Markdown
    md_content = _to_markdown(text, filename)

    # Save to corpus
    doc_id = _safe_doc_id(filename)
    out_path = corpus_dir / f"{doc_id}.md"

    # If the same doc already exists, overwrite it (re-upload = re-process)
    if out_path.exists() and doc_index:
        doc_index.remove_doc(doc_id)  # Clear old index data so it gets re-understood

    out_path.write_text(md_content, encoding="utf-8")

    # ── Pre-chunk + classify + tag + register in DocIndex ─────────────
    chunks = _split_into_chunks(md_content)
    chunk_meta_list = _classify_and_tag_chunks(chunks)

    if doc_index is not None:
        from antigravity.doc_index import ChunkMeta
        meta_objects = [
            ChunkMeta(
                chunk_id=m["chunk_id"],
                chunk_type=m["chunk_type"],
                mode=m["mode"],
                header=m["header"],
            )
            for m in chunk_meta_list
        ]
        doc_index.register_doc(
            doc_id=doc_id,
            filename=file_path.name,
            chunks=chunks,
            chunk_meta=meta_objects,
        )

    return {
        "doc_id": doc_id,
        "filename": file_path.name,
        "output_path": str(out_path),
        "chunks": len(chunks),
        "size": len(md_content),
    }


def _extract_pdf(path: Path) -> str:
    """Extract text from PDF using PyPDF2."""
    from PyPDF2 import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"## Page {i + 1}\n\n{text.strip()}")
    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    """Extract text from DOCX using python-docx."""
    from docx import Document

    doc = Document(str(path))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Preserve heading styles
        if para.style and para.style.name.startswith("Heading"):
            level = para.style.name.replace("Heading ", "").strip()
            try:
                hashes = "#" * int(level)
            except ValueError:
                hashes = "##"
            paragraphs.append(f"{hashes} {text}")
        else:
            paragraphs.append(text)

    # Also extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            header = rows[0]
            sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
            paragraphs.append("\n".join([header, sep] + rows[1:]))

    return "\n\n".join(paragraphs)


def _to_markdown(text: str, title: str) -> str:
    """Wrap extracted text in a clean Markdown document."""
    # Clean up excessive whitespace
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+\n", "\n", text)

    return f"# {title}\n\n{text.strip()}\n"


def _safe_doc_id(name: str) -> str:
    """Convert filename to a safe document ID."""
    safe = re.sub(r"[^a-zA-Z0-9_-]", "_", name)
    safe = re.sub(r"_+", "_", safe).strip("_")
    return safe.lower() if safe else "document"


def _split_into_chunks(content: str, max_chunk: int = 1500) -> list[str]:
    """Split document into chunks by headings or paragraph groups."""
    sections = re.split(r"\n(?=#+\s)", content)
    chunks: list[str] = []
    for section in sections:
        section = section.strip()
        if not section:
            continue
        if len(section) <= max_chunk:
            chunks.append(section)
        else:
            paras = section.split("\n\n")
            current = ""
            for para in paras:
                if len(current) + len(para) + 2 > max_chunk and current:
                    chunks.append(current.strip())
                    current = para
                else:
                    current += "\n\n" + para if current else para
            if current.strip():
                chunks.append(current.strip())
    return chunks if chunks else [content]


def _classify_and_tag_chunks(chunks: list[str]) -> list[dict]:
    """Classify each chunk and tag it with target model mode."""
    from antigravity.chunker import classify_chunk
    from antigravity.models import CHUNK_TYPE_TO_MODE

    result = []
    for i, chunk_text in enumerate(chunks):
        chunk_type = classify_chunk(chunk_text)
        mode = CHUNK_TYPE_TO_MODE[chunk_type]

        # Extract nearest heading as header
        header = ""
        for line in chunk_text.split("\n"):
            if line.startswith("#"):
                header = line.lstrip("#").strip()
                break

        result.append({
            "chunk_id": f"C{i}",
            "chunk_type": chunk_type.value,
            "mode": mode.value,
            "header": header,
        })

    return result

